"""
app.py – Opportunity Finder
Job-board style Streamlit app: search bar + category pill filters on the
home page, results list, and a detail view with AI motivation-letter
generation.
"""

import os
import io
import random

import streamlit as st
from docx import Document

from modules import profile_manager as pm
from modules import search_engine as se
from modules import ai_engine as ai

# ─────────────────────────────────────────────
# Page config
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Opportunity Finder",
    page_icon="🧪",
    layout="wide",
)

# ─────────────────────────────────────────────
# Global CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif;
}

/* ── Top brand bar ── */
.brand-row { display: flex; align-items: center; gap: 10px; margin-bottom: 2px; }
.brand-logo { font-size: 1.9rem; }
.brand-name { font-size: 1.6rem; font-weight: 800; color: #111827; }
.brand-tag { font-size: 0.92rem; color: #6b7280; margin: 2px 0 18px 0; }

/* ── Search bar block ── */
.search-block {
    background: #ffffff;
    border-radius: 16px;
    padding: 18px 20px 6px;
    border: 1.5px solid #eef0f4;
    box-shadow: 0 4px 16px rgba(0,0,0,0.05);
    margin-bottom: 14px;
}

/* Round every text input / selectbox like a real search bar */
div[data-testid="stTextInput"] input,
div[data-baseweb="select"] > div {
    border-radius: 10px !important;
}

/* Primary search button -> pill shaped, prominent */
div[data-testid="stButton"] button[kind="primary"] {
    border-radius: 10px;
    font-weight: 700;
    height: 42px;
}

/* Pill / chip filter buttons */
.pill-row div[data-testid="stButton"] button {
    border-radius: 999px !important;
    font-size: 0.85rem;
    font-weight: 600;
    padding: 6px 4px;
}
.pill-row div[data-testid="stButton"] button[kind="primary"] {
    background: #111827;
    border-color: #111827;
    color: #fff;
}
.pill-row div[data-testid="stButton"] button[kind="secondary"] {
    background: #f3f4f6;
    border: 1.5px solid #f3f4f6;
    color: #374151;
}

/* ── Opportunity list cards ── */
.opp-card {
    background: #ffffff;
    border-radius: 14px;
    padding: 20px 22px;
    margin-bottom: 10px;
    border: 1.5px solid #f0f0f5;
    box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    transition: box-shadow .15s;
}
.opp-card:hover { box-shadow: 0 6px 20px rgba(0,0,0,0.10); }
.opp-title { font-size: 1.05rem; font-weight: 700; color: #111827; }
.opp-org { font-size: 0.88rem; color: #6b7280; margin-top: 2px; }
.opp-meta { display: flex; flex-wrap: wrap; gap: 8px; margin-top: 10px; }

/* ── Badges ── */
.badge {
    display: inline-flex; align-items: center;
    padding: 3px 11px; border-radius: 999px;
    font-size: 0.73rem; font-weight: 600;
}
.b-country { background: #eff6ff; color: #1d4ed8; }
.b-field { background: #fff7ed; color: #c2410c; }
.b-fund { background: #f0fdf4; color: #166534; }
.b-degree { background: #faf5ff; color: #7e22ce; }
.b-deadline { background: #fef2f2; color: #b91c1c; }
.b-type { background: #ecfeff; color: #0e7490; }

/* ── Detail page ── */
.detail-header {
    background: linear-gradient(135deg, #1a3c5e 0%, #0d6eaa 100%);
    border-radius: 18px;
    padding: 32px 36px;
    color: white;
    margin-bottom: 24px;
}
.detail-title { font-size: 1.5rem; font-weight: 700; line-height: 1.3; }
.detail-org { font-size: 1rem; opacity: .85; margin-top: 6px; }

/* ── Section headings ── */
.section-label {
    font-size: 0.75rem; font-weight: 700; letter-spacing: .08em;
    text-transform: uppercase; color: #9ca3af; margin-bottom: 6px;
}

.results-count { color: #6b7280; font-size: 0.85rem; margin: 10px 0 12px; }

/* hide streamlit chrome */
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# Session state
# ─────────────────────────────────────────────
if "profile" not in st.session_state:
    st.session_state.profile = pm.load_profile()
if "opportunities" not in st.session_state:
    st.session_state.opportunities = pm.load_opportunities()
if "search_seed" not in st.session_state:
    st.session_state.search_seed = random.randint(0, 1_000_000)
if "api_key" not in st.session_state:
    st.session_state["api_key"] = os.environ.get("GROQ_API_KEY", "")

# Navigation state: "home" | "detail"
if "page" not in st.session_state:
    st.session_state.page = "home"
if "active_category" not in st.session_state:
    st.session_state.active_category = "all"
if "selected_opp" not in st.session_state:
    st.session_state.selected_opp = None

# ─────────────────────────────────────────────
# Categories (pill filters)
# ─────────────────────────────────────────────
ALL_PILL = {
    "key": "all",
    "icon": "🗂️",
    "title": "All",
    "desc": "Everything you've saved",
    "color": "#374151",
}

CATEGORIES = [
    {
        "key": "phd",
        "icon": "🎓",
        "title": "PhD",
        "desc": "Doctoral programmes & funded research posts",
        "color": "#4f46e5",
        "opp_type": "PhD Position",
        "field_hint": "chemical engineering PhD",
    },
    {
        "key": "industrial",
        "icon": "🏭",
        "title": "Industry",
        "desc": "Hands-on placements in industry & companies",
        "color": "#0891b2",
        "opp_type": "Internship",
        "field_hint": "chemical engineering industrial internship",
    },
    {
        "key": "research",
        "icon": "🔬",
        "title": "Research",
        "desc": "Lab & university research programmes",
        "color": "#059669",
        "opp_type": "Research Position",
        "field_hint": "chemical engineering research internship",
    },
    {
        "key": "other",
        "icon": "🌟",
        "title": "Other",
        "desc": "Fellowships, scholarships, summer programmes",
        "color": "#d97706",
        "opp_type": "Fellowship",
        "field_hint": "chemical engineering fellowship scholarship",
    },
]

PILLS = [ALL_PILL] + CATEGORIES
CAT_MAP = {c["key"]: c for c in PILLS}

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────
def text_to_docx_bytes(text: str, title: str = "") -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for line in text.split("\n"):
        if line.strip().startswith("Subject:"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif (
            line.strip().upper() == line.strip()
            and len(line.strip()) > 3
            and line.strip() != ""
        ):
            doc.add_heading(line.strip(), level=2)
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def get_category_opps(cat_key):
    """Filter the already-fetched opportunities by category (in-memory, no re-search)."""
    opps_all = list(reversed(st.session_state.opportunities))
    if cat_key == "all":
        return opps_all
    cat = CAT_MAP[cat_key]
    keyword = cat["opp_type"].lower()
    return [
        o for o in opps_all
        if keyword in o.get("title", "").lower() or keyword in o.get("summary", "").lower()
    ]


def nav_home():
    st.session_state.page = "home"
    st.session_state.selected_opp = None


def nav_detail(opp):
    st.session_state.selected_opp = opp
    st.session_state.page = "detail"


# ─────────────────────────────────────────────
# Sidebar (account / API key only — search lives on the home page now)
# ─────────────────────────────────────────────
with st.sidebar:
    st.header("🔑 Groq API Key")
    st.text_input(
        "Groq API Key",
        type="password",
        help="Free key at https://console.groq.com",
        key="api_key",
    )
    st.divider()
    st.header("👤 Your Profile")
    profile = st.session_state.profile
    profile["name"] = st.text_input("Full name", value=profile.get("name", ""))
    profile["summary"] = st.text_area(
        "Personal description / goals",
        value=profile.get("summary", ""),
        height=160,
        placeholder="e.g. 3rd-year Chemical Engineering student interested in sustainable polymers…",
    )
    if st.button("💾 Save profile", use_container_width=True):
        pm.save_profile(profile)
        st.success("Profile saved.")
    st.divider()
    st.caption(f"📊 {len(st.session_state.opportunities)} opportunities saved in total.")

# ═══════════════════════════════════════════════════════════
# PAGE: HOME  (search + category pills + results list)
# ═══════════════════════════════════════════════════════════
if st.session_state.page == "home":

    st.markdown(
        '<div class="brand-row"><span class="brand-logo">🧪</span>'
        '<span class="brand-name">Opportunity Finder</span></div>',
        unsafe_allow_html=True,
    )
    st.markdown(
        '<div class="brand-tag">Discover internships, PhD positions, research programmes, '
        'and fellowships — powered by live web search and AI.</div>',
        unsafe_allow_html=True,
    )

    # ── Search bar (front and center, like a real job board) ──
    with st.container():
        st.markdown('<div class="search-block">', unsafe_allow_html=True)
        s_col1, s_col2, s_col3 = st.columns([3, 2, 1.2])
        with s_col1:
            field = st.text_input(
                "Keywords",
                placeholder="🔍  Job title, field, or keywords (e.g. sustainable polymers)",
                key="search_field",
                label_visibility="collapsed",
            )
        with s_col2:
            place = st.text_input(
                "Location",
                placeholder="📍  Country / city / Anywhere",
                key="search_place",
                label_visibility="collapsed",
            )
        with s_col3:
            search_clicked = st.button("Search", type="primary", use_container_width=True)

        with st.expander("⚙️ Advanced filters"):
            f1, f2, f3, f4 = st.columns(4)
            with f1:
                degree = st.selectbox("Degree level", ["Any", "Undergraduate", "Master", "PhD", "Postdoc"], key="f_degree")
            with f2:
                funding = st.selectbox("Funding", ["Any", "Fully Funded", "Partially Funded", "Stipend", "Unfunded"], key="f_funding")
            with f3:
                duration = st.selectbox("Duration", ["Any", "1-2 months", "3 months", "6 months", "1 year", "2+ years"], key="f_duration")
            with f4:
                period = st.selectbox("Period", ["Any", "Summer 2026", "Fall 2026", "Spring 2027", "Rolling / Year-round"], key="f_period")
            include_social = st.checkbox("Include LinkedIn / social", value=True, key="f_social")
        st.markdown('</div>', unsafe_allow_html=True)

    if search_clicked:
        if not field.strip():
            st.warning("Enter a keyword or field to search for.")
        elif not st.session_state.get("api_key"):
            st.error("Enter your Groq API key in the sidebar first.")
        else:
            st.session_state.search_seed = random.randint(0, 1_000_000)
            with st.spinner("Searching the web…"):
                raw = se.search_opportunities(
                    field=field,
                    place=place or "Anywhere",
                    degree=degree,
                    funding=funding,
                    duration=duration,
                    period=period,
                    include_social=include_social,
                    seed=st.session_state.search_seed,
                )
            if not raw:
                st.error("No results. DuckDuckGo may be rate-limiting — wait 30s and retry.")
            else:
                st.info(f"Found {len(raw)} raw results — organising with AI…")
                structured = None
                with st.spinner("Organising results…"):
                    try:
                        structured = ai.structure_opportunities(
                            raw, field, api_key=st.session_state["api_key"]
                        )
                    except ValueError as e:
                        st.error(f"❌ API key error: {e}")
                    except Exception as e:
                        st.error(f"❌ Unexpected error while calling AI: {e}")

                if structured:
                    merged, added = pm.merge_opportunities(
                        st.session_state.opportunities, structured
                    )
                    st.session_state.opportunities = merged
                    pm.save_opportunities(merged)
                    st.session_state.active_category = "all"
                    st.success(f"✅ Added {added} new opportunit{'y' if added == 1 else 'ies'}.")
                    st.rerun()
                elif structured is not None:
                    st.warning(
                        "⚠️ AI returned no structured opportunities from these results. "
                        "This can happen when the LLM response isn't valid JSON. "
                        "Try pressing Search again — results rotate each time."
                    )

    st.divider()

    # ── Category pills (filter the already-fetched list, no re-search) ──
    counts = {c["key"]: len(get_category_opps(c["key"])) for c in PILLS}
    st.markdown('<div class="pill-row">', unsafe_allow_html=True)
    pill_cols = st.columns(len(PILLS))
    for col, cat in zip(pill_cols, PILLS):
        with col:
            is_active = st.session_state.active_category == cat["key"]
            label = f"{cat['icon']} {cat['title']} · {counts[cat['key']]}"
            if st.button(
                label,
                key=f"pill_{cat['key']}",
                use_container_width=True,
                type="primary" if is_active else "secondary",
            ):
                st.session_state.active_category = cat["key"]
                st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

    # ── Filter-within-results bar ──
    local_q = st.text_input(
        "", placeholder="🔎 Filter these results by title, organization, or location…",
        key="local_filter", label_visibility="collapsed",
    )

    active_cat = CAT_MAP[st.session_state.active_category]
    opps = get_category_opps(st.session_state.active_category)

    if local_q:
        q_lower = local_q.lower()
        opps = [
            o for o in opps
            if q_lower in o.get("title", "").lower()
            or q_lower in o.get("organization", "").lower()
            or q_lower in o.get("field", "").lower()
            or q_lower in o.get("location", "").lower()
        ]

    if not opps:
        if not st.session_state.opportunities:
            st.info("👋 No opportunities yet — use the search bar above to find your first ones!")
        else:
            st.info(f"No {active_cat['title'].lower()} opportunities match right now. Try another category or search again.")
    else:
        st.markdown(
            f"<div class='results-count'>{len(opps)} {active_cat['title'].lower()} opportunit{'y' if len(opps) == 1 else 'ies'} found</div>",
            unsafe_allow_html=True,
        )
        for idx, opp in enumerate(opps):
            title = opp.get("title", "Untitled")
            org = opp.get("organization", "Unknown")
            loc = opp.get("location", "Not specified")
            field_val = opp.get("field", "Not specified")
            funding_val = opp.get("funding", "")
            deadline_val = opp.get("deadline", "")
            degree_val = opp.get("degree_level", "")

            badges = ""
            if loc and loc != "Not specified":
                badges += f'<span class="badge b-country">📍 {loc}</span>'
            if field_val and field_val != "Not specified":
                badges += f'<span class="badge b-field">🔬 {field_val}</span>'
            if funding_val and funding_val != "Not specified":
                badges += f'<span class="badge b-fund">💰 {funding_val}</span>'
            if degree_val and degree_val not in ("Not specified", "Any"):
                badges += f'<span class="badge b-degree">🎓 {degree_val}</span>'
            if deadline_val and deadline_val != "Not specified":
                badges += f'<span class="badge b-deadline">⏰ {deadline_val}</span>'

            st.markdown(
                f"""
                <div class="opp-card">
                    <div class="opp-title">{title}</div>
                    <div class="opp-org">{org}</div>
                    <div class="opp-meta">{badges}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            if st.button("View details →", key=f"detail_{idx}", use_container_width=False):
                nav_detail(opp)
                st.rerun()
            st.markdown("<div style='height:4px'></div>", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════
# PAGE: DETAIL
# ═══════════════════════════════════════════════════════════
elif st.session_state.page == "detail":
    opp = st.session_state.selected_opp
    cat = CAT_MAP.get(st.session_state.active_category, ALL_PILL)

    col_back, _ = st.columns([1.2, 6])
    with col_back:
        if st.button("← Back to results"):
            nav_home()
            st.rerun()

    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)

    title = opp.get("title", "Untitled")
    org = opp.get("organization", "")
    st.markdown(
        f"""
        <div class="detail-header">
            <div class="detail-title">{title}</div>
            <div class="detail-org">{cat['icon']} {org}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        st.markdown('<div class="section-label">Location</div>', unsafe_allow_html=True)
        st.markdown(f"📍 {opp.get('location', 'Not specified')}")
        st.markdown('<div class="section-label" style="margin-top:16px">Field</div>', unsafe_allow_html=True)
        st.markdown(f"🔬 {opp.get('field', 'Not specified')}")
    with c2:
        st.markdown('<div class="section-label">Funding</div>', unsafe_allow_html=True)
        st.markdown(f"💰 {opp.get('funding', 'Not specified')}")
        st.markdown('<div class="section-label" style="margin-top:16px">Degree Level</div>', unsafe_allow_html=True)
        st.markdown(f"🎓 {opp.get('degree_level', 'Not specified')}")
    with c3:
        st.markdown('<div class="section-label">Deadline</div>', unsafe_allow_html=True)
        st.markdown(f"⏰ {opp.get('deadline', 'Not specified')}")
        st.markdown('<div class="section-label" style="margin-top:16px">Duration</div>', unsafe_allow_html=True)
        st.markdown(f"⌛ {opp.get('duration', 'Not specified')}")

    st.divider()

    st.subheader("📋 About this Opportunity")
    summary = opp.get("summary", "No description available.")
    st.markdown(summary)

    extra_fields = {
        "Period": opp.get("period"),
        "Contact": opp.get("contact"),
        "Source": opp.get("source"),
    }
    for label, val in extra_fields.items():
        if val and val not in ("Not specified", None, ""):
            st.markdown(f"**{label}:** {val}")

    link = opp.get("link", "")
    if link and link != "#":
        st.link_button("🔗 Open original listing", link)

    st.divider()

    st.subheader("✨ Generate Motivation Letter")
    profile = st.session_state.profile
    has_profile = bool(profile.get("name") or profile.get("summary"))

    if not has_profile:
        st.warning("Fill in your name and personal description in the sidebar so the AI can personalise your letter.")
    else:
        if st.button("📝 Generate Motivation Letter", type="primary"):
            if not st.session_state.get("api_key"):
                st.error("❌ Enter your Groq API key in the sidebar first.")
            else:
                profile_text = pm.get_profile_text(profile)
                with st.spinner("Writing your letter…"):
                    try:
                        letter = ai.generate_motivation_letter(
                            profile_text, opp, api_key=st.session_state["api_key"]
                        )
                        st.session_state["detail_letter"] = letter
                    except RuntimeError as e:
                        st.error(f"❌ {e}")
                    except Exception as e:
                        st.error(f"❌ Unexpected error: {e}")

        if "detail_letter" in st.session_state:
            st.text_area("Motivation Letter", st.session_state["detail_letter"], height=400)
            st.download_button(
                "⬇️ Download as .docx",
                data=text_to_docx_bytes(st.session_state["detail_letter"], "Motivation Letter"),
                file_name="motivation_letter.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    st.divider()
    st.caption("⚠️ Always verify deadlines on the original listing. Review and personalise AI-generated text before sending.")

